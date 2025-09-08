# document_parser.py
import os
import io
import base64
import hashlib
from typing import Dict, List, Any

from docx import Document
import PyPDF2
import openpyxl
from PIL import Image

def _resolve(obj):
    """PyPDF2 IndirectObject 등을 실제 객체로 해제"""
    try:
        while hasattr(obj, "get_object"):
            obj = obj.get_object()
    except Exception:
        pass
    return obj

def _as_float_list(x):
    x = _resolve(x)
    if x is None:
        return None
    try:
        # ArrayObject 등도 list(...) 가능
        seq = list(x) if not isinstance(x, (list, tuple)) else x
        out = []
        for v in seq:
            try:
                out.append(float(_resolve(v)))
            except Exception:
                # 좌표가 숫자 외 타입이면 스킵
                pass
        return out if out else None
    except Exception:
        return None

def _safe_text(v):
    v = _resolve(v)
    if v is None:
        return ""
    try:
        return str(v)
    except Exception:
        try:
            return v.decode("utf-8", "replace")
        except Exception:
            return repr(v)
def _float_list(x):
    if x is None:
        return None
    try:
        return [float(v) for v in x]
    except Exception:
        return None


def _name_value(v):
    """PyPDF2 NameObject -> 'Text' 같은 문자열로 정규화"""
    try:
        return v.name.lstrip('/')  # NameObject
    except Exception:
        return str(v).lstrip('/')


def _safe_str(v):
    return "" if v is None else str(v)


def _compute_annot_fallback_id(page_number: int, subtype: str, rect: List[float], contents: str) -> str:
    base = f"{page_number}:{subtype}:{','.join(f'{n:.2f}' for n in (rect or []))}:{contents}"
    digest = hashlib.md5(base.encode('utf-8')).hexdigest()[:10]
    return f"AUTO-{page_number}-{subtype}-{digest}"


class DocumentParser:
    """
    DOCX / PDF / XLSX 문서에서 비교에 필요한 구조화 데이터를 추출한다.
    - DOCX: paragraphs (runs 포함 포맷), tables, images(base64 PNG)
    - PDF : pages(text, rotation, mediabox), annotations(페이지 평탄화)
    - XLSX: sheets(cells with coordinate/value + font/fill/border)
    """
    def __init__(self):
        self.supported_formats = {
            '.docx': self._parse_docx,
            '.pdf':  self._parse_pdf,
            '.xlsx': self._parse_xlsx,
        }

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        return self.supported_formats[file_ext](file_path)

    # ---------------------------------------------------------------------
    # DOCX
    # ---------------------------------------------------------------------
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        content: Dict[str, Any] = {
            'type': 'docx',
            'text_content': [],
            'tables': [],
            'images': [],
            'formatting': [],
            'paragraphs': [],
        }

        # Paragraphs + run-level formatting
        for i, paragraph in enumerate(doc.paragraphs):
            para_data = {
                'index': i,
                'text': paragraph.text,
                'style': paragraph.style.name if paragraph.style else None,
                'runs': []
            }
            for run in paragraph.runs:
                run_data = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                }
                para_data['runs'].append(run_data)

            content['paragraphs'].append(para_data)
            content['text_content'].append(paragraph.text)

        # Tables
        for i, table in enumerate(doc.tables):
            table_data = {
                'index': i,
                'rows': [],
                'style': table.style.name if table.style else None
            }
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append({
                        'text': cell.text,
                        'paragraphs': [p.text for p in cell.paragraphs],
                    })
                table_data['rows'].append(row_data)
            content['tables'].append(table_data)

        # Images
        for i, rel in enumerate(doc.part.rels.values()):
            if getattr(rel, "reltype", "").endswith("/image") or "image" in getattr(rel, "target_ref", ""):
                try:
                    blob = rel.target_part.blob
                    img = Image.open(io.BytesIO(blob))
                    buf = io.BytesIO()
                    img.save(buf, format='PNG')
                    img_b64 = base64.b64encode(buf.getvalue()).decode()

                    content['images'].append({
                        'index': i,
                        'data': img_b64,
                        'format': img.format,
                        'size': img.size,
                        'mode': img.mode,
                    })
                except Exception as e:
                    print(f"[DOCX] Error processing image {i}: {e}")

        return content

    # ---------------------------------------------------------------------
    # PDF (annotations 포함)
    # ---------------------------------------------------------------------
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        content: Dict[str, Any] = {
            'type': 'pdf',
            'text_content': [],
            'pages': [],
            'images': [],
            'annotations': [],
        }

        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            for page_num, page in enumerate(reader.pages):
                # rotation
                try:
                    rotation = page.get('/Rotate', 0)
                except Exception:
                    rotation = 0

                # mediabox → 직렬화 가능한 리스트 형태로
                mb = getattr(page, "mediabox", None)
                if mb is not None:
                    try:
                        mediabox = [float(mb.left), float(mb.bottom), float(mb.right), float(mb.top)]
                    except Exception:
                        mediabox = [0, 0, 0, 0]
                else:
                    mediabox = [0, 0, 0, 0]

                text = page.extract_text() or ""
                content['pages'].append({
                    'page_number': page_num + 1,
                    'text': text,
                    'rotation': rotation,
                    'mediabox': mediabox,
                })
                content['text_content'].append(text)

                # -------- Annotations (여기가 핵심 수정) --------
                annots_ref = page.get('/Annots', None)
                if annots_ref is None:
                    # 어떤 PDF는 dict 접근으로만 되는 경우가 있어 보조 시도
                    try:
                        annots_ref = page['/Annots']
                    except Exception:
                        annots_ref = None

                if not annots_ref:
                    continue

                annots_arr = _resolve(annots_ref)  # IndirectObject → ArrayObject 로 해제
                try:
                    iterable = list(annots_arr)
                except Exception:
                    # 여전히 iterable 아니면 스킵
                    continue

                for i, annot_ref in enumerate(iterable):
                    try:
                        obj = _resolve(annot_ref)  # 각 주석 dict 로 해제
                        # dict 유사체만 처리
                        if not hasattr(obj, "get"):
                            continue

                        subtype = _safe_text(obj.get('/Subtype')).lstrip('/')
                        rect = _as_float_list(obj.get('/Rect'))
                        quad = _as_float_list(obj.get('/QuadPoints'))
                        contents = _safe_text(obj.get('/Contents'))
                        author = _safe_text(obj.get('/T'))
                        subject = _safe_text(obj.get('/Subj'))
                        color = _as_float_list(obj.get('/C'))
                        flags = _resolve(obj.get('/F'))
                        mod_date = _safe_text(obj.get('/M'))
                        creation = _safe_text(obj.get('/CreationDate'))
                        nm_raw = obj.get('/NM')
                        nm = _safe_text(nm_raw)

                        if not nm:
                            nm = _compute_annot_fallback_id(page_num + 1, subtype, rect or [], contents)

                        content['annotations'].append({
                            'id': nm,
                            'page_number': page_num + 1,
                            'subtype': subtype,  # Text, Highlight, StrikeOut, Underline, FreeText, ...
                            'rect': rect,  # [x1, y1, x2, y2]
                            'quadpoints': quad,  # 텍스트 마크업 주석의 경우 영역
                            'contents': contents,  # 코멘트/본문
                            'author': author,
                            'subject': subject,
                            'color': color,  # [r,g,b] 0..1
                            'flags': int(flags) if isinstance(flags, (int, float)) else None,
                            'modified': mod_date,
                            'created': creation,
                        })
                    except Exception as e:
                        print(f"[PDF] Annotation parse error on page {page_num + 1}: {e}")

        return content

    # ---------------------------------------------------------------------
    # XLSX
    # ---------------------------------------------------------------------
    def _parse_xlsx(self, file_path: str) -> Dict[str, Any]:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        content: Dict[str, Any] = {
            'type': 'xlsx',
            'sheets': [],
            'tables': []
        }

        for sheet_name in wb.sheetnames:
            sh = wb[sheet_name]
            sheet_data = {
                'name': sheet_name,
                'max_row': sh.max_row,
                'max_column': sh.max_column,
                'cells': []
            }
            for row in sh.iter_rows():
                row_out: List[Dict[str, Any]] = []
                for cell in row:
                    if cell.value is None:
                        continue
                    cell_dict: Dict[str, Any] = {
                        'coordinate': cell.coordinate,
                        'value': cell.value,
                        'data_type': type(cell.value).__name__,
                    }
                    if cell.font:
                        cell_dict['font'] = {
                            'name': cell.font.name,
                            'size': cell.font.size,
                            'bold': cell.font.bold,
                            'italic': cell.font.italic,
                            'color': str(cell.font.color.rgb) if (cell.font.color and cell.font.color.rgb) else None,
                        }
                    else:
                        cell_dict['font'] = None
                    if cell.fill:
                        cell_dict['fill'] = {
                            'pattern_type': cell.fill.patternType,
                            'fg_color': str(cell.fill.fgColor.rgb) if (cell.fill.fgColor and cell.fill.fgColor.rgb) else None,
                        }
                    else:
                        cell_dict['fill'] = None
                    if cell.border:
                        cell_dict['border'] = {
                            'left':   str(cell.border.left.style) if cell.border.left and cell.border.left.style else None,
                            'right':  str(cell.border.right.style) if cell.border.right and cell.border.right.style else None,
                            'top':    str(cell.border.top.style) if cell.border.top and cell.border.top.style else None,
                            'bottom': str(cell.border.bottom.style) if cell.border.bottom and cell.border.bottom.style else None,
                        }
                    else:
                        cell_dict['border'] = None
                    row_out.append(cell_dict)
                if row_out:
                    sheet_data['cells'].append(row_out)
            content['sheets'].append(sheet_data)

        return content
