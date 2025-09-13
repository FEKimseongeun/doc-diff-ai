# app/services/extractors.py
import os
from typing import Dict, List, Any, Optional
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook


class ExcelExtractor:
    """Excel 파일 내용 추출기"""

    def extract(self, file_path: str) -> Dict[str, List[List]]:
        """
        Excel 파일에서 데이터 추출

        Returns:
            {sheet_name: [[cell_values], ...]}
        """
        result = {}

        # pandas를 사용한 방법 (더 안정적)
        try:
            excel_file = pd.ExcelFile(file_path)

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

                # NaN 값을 None으로 변환
                df = df.where(pd.notnull(df), None)

                # 리스트로 변환
                sheet_data = df.values.tolist()
                result[sheet_name] = sheet_data

        except Exception as e:
            # openpyxl을 사용한 대체 방법
            wb = load_workbook(file_path, data_only=True)

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_data = []

                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        value = cell.value
                        # 날짜 처리
                        if hasattr(value, 'strftime'):
                            value = value.strftime('%Y-%m-%d %H:%M:%S')
                        row_data.append(value)
                    sheet_data.append(row_data)

                result[sheet_name] = sheet_data

            wb.close()

        return result

    def extract_with_formatting(self, file_path: str) -> Dict[str, Dict]:
        """
        서식 정보를 포함한 Excel 데이터 추출

        Returns:
            {sheet_name: {
                'data': [[cell_values], ...],
                'formatting': {...}
            }}
        """
        result = {}
        wb = load_workbook(file_path, data_only=True)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = []
            formatting = []

            for row_idx, row in enumerate(sheet.iter_rows()):
                row_data = []
                row_format = []

                for col_idx, cell in enumerate(row):
                    value = cell.value

                    # 날짜 처리
                    if hasattr(value, 'strftime'):
                        value = value.strftime('%Y-%m-%d %H:%M:%S')

                    row_data.append(value)

                    # 서식 정보 추출
                    cell_format = {
                        'font': {
                            'name': cell.font.name if cell.font else None,
                            'size': cell.font.size if cell.font else None,
                            'bold': cell.font.bold if cell.font else False,
                            'italic': cell.font.italic if cell.font else False,
                            'color': cell.font.color.rgb if cell.font and cell.font.color else None
                        },
                        'fill': {
                            'color': cell.fill.fgColor.rgb if cell.fill and cell.fill.fgColor else None
                        },
                        'alignment': {
                            'horizontal': cell.alignment.horizontal if cell.alignment else None,
                            'vertical': cell.alignment.vertical if cell.alignment else None
                        }
                    }
                    row_format.append(cell_format)

                sheet_data.append(row_data)
                formatting.append(row_format)

            result[sheet_name] = {
                'data': sheet_data,
                'formatting': formatting
            }

        wb.close()
        return result


class WordExtractor:
    """Word 문서 내용 추출기"""

    def extract(self, file_path: str) -> List[str]:
        """
        Word 문서에서 텍스트 추출

        Returns:
            [paragraph1, paragraph2, ...]
        """
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 표의 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))

        return paragraphs

    def extract_with_formatting(self, file_path: str) -> Dict[str, Any]:
        """
        서식 정보를 포함한 Word 문서 추출

        Returns:
            {
                'paragraphs': [...],
                'tables': [...],
                'formatting': {...}
            }
        """
        doc = Document(file_path)
        result = {
            'paragraphs': [],
            'tables': [],
            'formatting': []
        }

        # 단락 추출
        for para in doc.paragraphs:
            if para.text.strip():
                para_data = {
                    'text': para.text,
                    'style': para.style.name if para.style else None,
                    'alignment': str(para.alignment) if para.alignment else None,
                    'runs': []
                }

                # Run 단위 서식 정보
                for run in para.runs:
                    run_data = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'color': run.font.color.rgb if run.font.color else None
                    }
                    para_data['runs'].append(run_data)

                result['paragraphs'].append(para_data)

        # 표 추출
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'index': table_idx,
                'rows': []
            }

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_data = {
                        'text': cell.text.strip(),
                        'paragraphs': []
                    }

                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_data['paragraphs'].append({
                                'text': para.text,
                                'style': para.style.name if para.style else None
                            })

                    row_data.append(cell_data)

                table_data['rows'].append(row_data)

            result['tables'].append(table_data)

        return result


class PDFExtractor:
    """PDF 문서 내용 추출기"""

    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        PDF 문서에서 텍스트 추출

        Returns:
            {
                'pages': [page1_text, page2_text, ...],
                'metadata': {...}
            }
        """
        doc = fitz.open(file_path)
        result = {
            'pages': [],
            'metadata': {
                'page_count': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'keywords': doc.metadata.get('keywords', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creationDate': str(doc.metadata.get('creationDate', '')),
                'modDate': str(doc.metadata.get('modDate', ''))
            }
        }

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            result['pages'].append(text)

        doc.close()
        return result

    def extract_with_layout(self, file_path: str) -> Dict[str, Any]:
        """
        레이아웃 정보를 포함한 PDF 추출

        Returns:
            {
                'pages': [...],
                'metadata': {...}
            }
        """
        doc = fitz.open(file_path)
        result = {
            'pages': [],
            'metadata': {
                'page_count': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', '')
            }
        }

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # 텍스트 블록 단위로 추출
            blocks = page.get_text("blocks")

            page_data = {
                'page_number': page_num + 1,
                'width': page.rect.width,
                'height': page.rect.height,
                'blocks': []
            }

            for block in blocks:
                x0, y0, x1, y1, text, block_no, block_type = block

                block_data = {
                    'bbox': [x0, y0, x1, y1],
                    'text': text.strip(),
                    'block_no': block_no,
                    'type': 'text' if block_type == 0 else 'image'
                }

                if block_data['text']:  # 빈 텍스트 제외
                    page_data['blocks'].append(block_data)

            # 이미지 정보 추출
            image_list = page.get_images()
            page_data['images'] = []

            for img_idx, img in enumerate(image_list):
                img_data = {
                    'index': img_idx,
                    'xref': img[0],
                    'width': img[2],
                    'height': img[3]
                }
                page_data['images'].append(img_data)

            result['pages'].append(page_data)

        doc.close()
        return result

    def extract_tables(self, file_path: str) -> List[Dict]:
        """PDF에서 표 추출 (실험적)"""
        import camelot

        tables = []
        try:
            # Camelot을 사용한 표 추출
            table_list = camelot.read_pdf(file_path, pages='all')

            for idx, table in enumerate(table_list):
                tables.append({
                    'index': idx,
                    'page': table.page,
                    'data': table.df.values.tolist(),
                    'accuracy': table.accuracy
                })
        except Exception as e:
            print(f"표 추출 실패: {e}")
            # 대체 방법: PyMuPDF의 표 감지
            tables = self._extract_tables_pymupdf(file_path)

        return tables

    def _extract_tables_pymupdf(self, file_path: str) -> List[Dict]:
        """PyMuPDF를 사용한 표 추출"""
        doc = fitz.open(file_path)
        tables = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # 간단한 표 감지 로직
            # 실제로는 더 복잡한 알고리즘이 필요
            text = page.get_text("dict")

            # 표 패턴 감지 (예: 정렬된 텍스트 블록)
            # 이 부분은 실제 구현 시 개선 필요

        doc.close()
        return tables